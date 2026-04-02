import logging
import os
from flask import Flask, send_file, request, jsonify
from io import BytesIO
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from textblob import Word
from googletrans import Translator
import openai
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__, static_url_path='', static_folder='static')

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@app.route('/')
def root():
    logger.info(f"Request: {request.method} {request.path}")
    return app.send_static_file('index.html')

# -------------------------
# AUTOCORRECT ENDPOINT
# -------------------------
@app.post('/autocorrect')
def autocorrect():
    logger.info(f"Request: {request.method} {request.path}")
    data = request.get_json(force=True)
    word = data.get('word', '')
    if not word:
        return jsonify({"corrected": word})

    corrected = Word(word).correct()
    return jsonify({"corrected": str(corrected)})

# -------------------------
# TRANSLATE ENDPOINT
# -------------------------
@app.post('/translate')
def translate():
    logger.info(f"Request: {request.method} {request.path}")
    data = request.get_json(force=True)
    text = data.get('text', '')
    target = data.get('target', 'en')
    if not text:
        return jsonify({"translated": ""})

    translator = Translator()
    translated = translator.translate(text, dest=target)
    return jsonify({"translated": translated.text})

# -------------------------
# EXPORT TO DOCX
# -------------------------
@app.post('/export/docx')
def export_docx():
    logger.info(f"Request: {request.method} {request.path}")
    data = request.get_json(force=True)
    text = data.get('text', '')
    generated = data.get('generated', [])
    doc = Document()
    doc.add_heading('Assignment Draft', level=1)
    for para in text.split('\n'):
        doc.add_paragraph(para)
    if generated:
        doc.add_heading('Generated Content', level=2)
        for item in generated:
            if item['type'] == 'graph':
                doc.add_paragraph(f"Generated Graph: {item['description']}")
            elif item['type'] == 'equation':
                doc.add_paragraph(f"Generated Equation: {item['latex']}")
            elif item['type'] == 'diagram':
                doc.add_paragraph(f"Generated Diagram: {item['mermaid']}")
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name='assignment.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# -------------------------
# EXPORT TO PDF
# -------------------------
@app.post('/export/pdf')
def export_pdf():
    logger.info(f"Request: {request.method} {request.path}")
    data = request.get_json(force=True)
    text = data.get('text', '')
    generated = data.get('generated', [])
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    x, y = 50, height - 50
    c.setFont("Helvetica", 12)
    c.drawString(x, y, "Assignment Draft")
    y -= 20
    for line in text.split('\n'):
        if y < 50:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 50
        c.drawString(x, y, line)
        y -= 16
    if generated:
        if y < 100:
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 50
        c.drawString(x, y, "Generated Content")
        y -= 20
        for item in generated:
            if item['type'] == 'graph':
                c.drawString(x, y, f"Generated Graph: {item['description']}")
            elif item['type'] == 'equation':
                c.drawString(x, y, f"Generated Equation: {item['latex']}")
            elif item['type'] == 'diagram':
                c.drawString(x, y, f"Generated Diagram: {item['mermaid']}")
            y -= 16
            if y < 50:
                c.showPage()
                c.setFont("Helvetica", 12)
                y = height - 50
    c.save()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name='assignment.pdf',
                     mimetype='application/pdf')

# -------------------------
# GENERATE GRAPH ENDPOINT
# -------------------------
@app.post('/generate/graph')
def generate_graph():
    logger.info(f"Request: {request.method} {request.path}")
    data = request.get_json(force=True)
    description = data.get('description', '')
    if not description:
        return jsonify({"error": "Description is required"}), 400

    openai.api_key = os.getenv('OPENAI_API_KEY')
    if not openai.api_key or openai.api_key == 'your_openai_api_key_here':
        # Mock response for testing without API key
        mock_chart = {
            "type": "bar",
            "data": {
                "labels": ["January", "February", "March"],
                "datasets": [{
                    "label": "Sales",
                    "data": [10, 20, 30],
                    "backgroundColor": "rgba(75, 192, 192, 0.2)",
                    "borderColor": "rgba(75, 192, 192, 1)",
                    "borderWidth": 1
                }]
            },
            "options": {
                "scales": {
                    "y": {
                        "beginAtZero": True
                    }
                }
            }
        }
        return jsonify({"chart": str(mock_chart).replace("'", '"')})
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that generates JSON data for charts based on natural language descriptions. Output only valid JSON for Chart.js."},
                {"role": "user", "content": f"Generate Chart.js JSON data for: {description}. Include type (e.g., 'bar', 'line'), data, labels, and options."}
            ]
        )
        chart_json = response.choices[0].message.content.strip()
        return jsonify({"chart": chart_json})
    except Exception as e:
        logger.error(f"Error generating graph: {e}")
        return jsonify({"error": "Failed to generate graph"}), 500

# -------------------------
# GENERATE EQUATION ENDPOINT
# -------------------------
@app.post('/generate/equation')
def generate_equation():
    logger.info(f"Request: {request.method} {request.path}")
    data = request.get_json(force=True)
    description = data.get('description', '')
    if not description:
        return jsonify({"error": "Description is required"}), 400

    openai.api_key = os.getenv('OPENAI_API_KEY')
    if not openai.api_key or openai.api_key == 'your_openai_api_key_here':
        # Mock response for testing without API key
        return jsonify({"latex": "\\frac{a^2 + b^2}{c}"})
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a helpful assistant that generates LaTeX math equations based on natural language descriptions. Output only the LaTeX code."},
                {"role": "user", "content": f"Generate LaTeX code for the equation: {description}"}
            ]
        )
        latex = response.choices[0].message.content.strip()
        return jsonify({"latex": latex})
    except Exception as e:
        logger.error(f"Error generating equation: {e}")
        return jsonify({"error": "Failed to generate equation"}), 500

# -------------------------
# GENERATE DIAGRAM ENDPOINT
# -------------------------
@app.post('/generate/diagram')
def generate_diagram():
    logger.info(f"Request: {request.method} {request.path}")
    data = request.get_json(force=True)
    description = data.get('description', '')
    if not description:
        return jsonify({"error": "Description is required"}), 400

    openai.api_key = os.getenv('OPENAI_API_KEY')
    if not openai.api_key or openai.api_key == 'your_openai_api_key_here':
        # Mock response for testing without API key
        return jsonify({"mermaid": "graph TD\nA[Start] --> B[Process]\nB --> C[End]"})
    else:
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that generates Mermaid diagram code based on natural language descriptions. Output only valid Mermaid syntax."},
                    {"role": "user", "content": f"Generate Mermaid code for: {description}"}
                ]
            )
            mermaid = response.choices[0].message.content.strip()
            return jsonify({"mermaid": mermaid})
        except Exception as e:
            logger.error(f"Error generating diagram: {e}")
            return jsonify({"error": "Failed to generate diagram"}), 500

if __name__ == '__main__':
     app.run(host="0.0.0.0", port=5000, debug=True)
